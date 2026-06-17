"""
utils/output_format.py — 报表输出函数

提供三线表 docx 和 Excel 清单输出。
"""

import os
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn


def save_table_to_docx_threeline(
    df: pd.DataFrame,
    filepath: str,
    title: str = "",
    notes: list[str] | None = None,
    row_height_cm: float = 0.6,
    auto_width: bool = True,
):
    """将 DataFrame 输出为三线表 docx。"""
    os.makedirs(os.path.dirname(filepath), exist_ok=True)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    if title:
        p = doc.add_paragraph()
        p.alignment = 1
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(12)

    table = doc.add_table(rows=1 + len(df), cols=len(df.columns))
    table.style = "Table Grid"

    # 表头
    for j, col in enumerate(df.columns):
        cell = table.rows[0].cells[j]
        cell.text = str(col)

    # 数据行
    for i in range(len(df)):
        for j, col in enumerate(df.columns):
            cell = table.rows[i + 1].cells[j]
            val = df.iloc[i, j]
            cell.text = "" if pd.isna(val) else str(val)

    if notes:
        for note in notes:
            doc.add_paragraph(note, style="Normal")

    doc.save(filepath)


def export_to_excel_with_format(
    df: pd.DataFrame,
    filepath: str,
    sheet_name: str = "Sheet1",
    header_title: str | None = None,
):
    """将 DataFrame 输出为格式化的 Excel 清单。"""
    os.makedirs(os.path.dirname(filepath), exist_ok=True)

    with pd.ExcelWriter(filepath, engine="xlsxwriter") as writer:
        start_row = 1 if header_title else 0
        df.to_excel(writer, sheet_name=sheet_name, index=False, startrow=start_row)

        ws = writer.sheets[sheet_name]

        if header_title:
            ws.write(0, 0, header_title)

        # 自动列宽
        for j, col in enumerate(df.columns):
            max_len = max(
                df[col].astype(str).map(len).max() if len(df) > 0 else 0,
                len(str(col)),
            )
            ws.set_column(j, j, min(max_len + 2, 50))
